"""
Скрипт для парсинга педагогического состава из DOCX и загрузки в БД
"""
import sys
sys.path.append('backend')

from docx import Document
from backend.app.core.database import SessionLocal
from backend.app.models.staff import Staff

def parse_staff_docx(docx_path: str):
    """Парсит DOCX с педагогическим составом"""
    print(f"Открываем файл: {docx_path}")

    doc = Document(docx_path)
    print(f"Параграфов в документе: {len(doc.paragraphs)}")

    staff_data = []

    # Извлекаем текст
    print("\nПервые 30 строк документа:")
    for i, para in enumerate(doc.paragraphs[:30]):
        text = para.text.strip()
        if text:
            print(f"{i}: {text}")
            staff_data.append(text)

    # Проверяем таблицы
    if doc.tables:
        print(f"\n\nНайдено таблиц: {len(doc.tables)}")

        for table_num, table in enumerate(doc.tables[:2]):  # Первые 2 таблицы
            print(f"\n=== Таблица {table_num + 1} ===")
            print(f"Строк: {len(table.rows)}, Колонок: {len(table.columns)}")

            # Показываем первые строки
            for i, row in enumerate(table.rows[:5]):
                cells = [cell.text.strip() for cell in row.cells]
                print(f"Строка {i}: {cells}")

    return staff_data

def load_staff_to_db(staff_data):
    """Загружает данные о сотрудниках в БД"""
    db = SessionLocal()

    try:
        # Пример: создаем тестовые записи
        # TODO: Адаптировать под реальную структуру из DOCX

        test_staff = Staff(
            full_name="Иванов Иван Иванович",
            position="Преподаватель",
            department="Кафедра информационных технологий",
            room_number="305",
            email="ivanov@kkrit.ru",
            phone="+7 (123) 456-78-90"
        )

        db.add(test_staff)
        db.commit()
        print("✓ Тестовая запись добавлена в БД")

    except Exception as e:
        print(f"Ошибка при загрузке в БД: {e}")
        db.rollback()
    finally:
        db.close()

if __name__ == "__main__":
    docx_path = "ped._sostav_dlya_sayta_10.09.2025.docx"

    # Парсим DOCX
    staff_data = parse_staff_docx(docx_path)

    if staff_data:
        print("\n" + "="*50)
        print("DOCX успешно распарсен!")
        print("="*50)

        # Загружаем в БД
        # load_staff_to_db(staff_data)
        print("\nДля загрузки в БД нужно адаптировать структуру под реальные данные")
